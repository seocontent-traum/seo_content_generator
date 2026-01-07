# import torch
# import streamlit as st
# import requests
# from bs4 import BeautifulSoup
# from langchain.schema import Document
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from langchain_community.vectorstores import FAISS
# from sentence_transformers import SentenceTransformer
# # from langchain_openai import AzureChatOpenAI
# from docx import Document as DC
# from langchain_huggingface import HuggingFaceEmbeddings
# import google.generativeai as genai


# # Streamlit setup
# st.set_page_config(page_title="🧠 SEO Content Generator", page_icon="🤓", layout="wide")

# st.markdown(
#     """
#     <h1 style='text-align: center; color: #4CAF50;'>🧠 SEO Content Generator v2.0</h1>
#     <p style='text-align: center; font-size: 18px;'>Generate top-ranking SEO content. Powered by Gemini Pro & Google Search!</p>
#     <hr style="border:1px solid #ccc;">
#     """,
#     unsafe_allow_html=True
# )
# col1, col2 = st.columns([6, 1])
# with col2:
#    st.markdown(
#     """
#     <a href="https://www.linkedin.com/in/abhishek-kumawat-iitd/" target="_blank">
#         <button style="background-color:#0072b1; color:white; padding:8px 16px; border:none; border-radius:6px; cursor:pointer;">
#             💡 How it works?
#         </button>
#     </a>
#     """,
#     unsafe_allow_html=True
# )


# # Sidebar inputs
# st.sidebar.header("🛠️ Input Parameters")
# doc_file = "seo_content_revamp_AK.docx"

# # --- Language and Country Mapping ---
# language_options = {
#     "Danish": "lang_da",
#     "Dutch": "lang_nl",
#     "German": "lang_de",
#     "French": "lang_fr",
#     "English": "lang_en",
#     "Swedish": "lang_sv",
#     "Norwegian": "lang_no",
#     "Spanish": "lang_es",
# }

# country_options = {
#     "India": "in",
#     "Denmark": "dk",
#     "Netherlands": "nl",
#     "Belgium": "be",
#     "Germany": "de",
#     "France": "fr",
#     "United States": "us",
#     "Sweden": "se",
#     "Norway": "no",
#     "Spain": "es"
    
# }

# # --- Dropdowns in Sidebar ---
# language_choice = st.sidebar.selectbox("🌐 Language", list(language_options.keys()))
# country_choice = st.sidebar.selectbox("🏳️ Country", list(country_options.keys()))

# # --- Get API codes based on dropdown selections ---
# search_language = language_options[language_choice]
# search_country = country_options[country_choice]

# language = language_choice
# country = country_choice
# brand = st.sidebar.text_input("🏷️ Brand", "Danland")
# base_urls = st.sidebar.text_area("🔗 Base URLs (comma-separated)").split(",")
# keyword = st.sidebar.text_area("🔍 Enter the Keyword (Required)")
# additional_input = st.sidebar.text_area("✏️ Additional Input (Enter additional input here)")




# # Configure the Gemini client
# genai.configure(api_key=st.secrets["gemini_api"])

# # Initialize the model
# model = genai.GenerativeModel("gemini-2.5-pro")


# # Always load on CPU
# # device = 'cpu'

# # encoder = SentenceTransformer("all-mpnet-base-v2")

# # # Move empty weights to CPU, then load actual weights
# # encoder = encoder.to_empty(device=device)
# # encoder.load_state_dict(encoder.state_dict())
# # # encoder = encoder.to(device)

# # embeddings = HuggingFaceEmbeddings(
# #     model_name='sentence-transformers/all-MiniLM-L6-v2',
# #     model_kwargs={"device": device}
# # )

# # Always load on CPU
# device = 'cpu'

# # ✅ Load directly on CPU, avoiding .to() or .to_empty()
# encoder = SentenceTransformer("all-mpnet-base-v2", device=device)

# embeddings = HuggingFaceEmbeddings(
#     model_name='sentence-transformers/all-MiniLM-L6-v2',
#     model_kwargs={"device": device}
# )



# def extract_text_from_url(url):
#     try:
#         response = requests.get(url, timeout=5)
#         response.raise_for_status()
#         soup = BeautifulSoup(response.content, 'html.parser')
#         return " ".join([p.get_text() for p in soup.find_all("p")])
#     except requests.exceptions.RequestException as e:
#         st.warning(f"Error fetching {url}: {e}")
#         return ""

# def generate_keyword(text):
#     query = f"Based on the {text}. Tell me a keyword which sums up all content."
#     response = model.generate_content(query)
#     return response.text 

# # ✅ NEW: Google Custom Search API
# def scrape_google_search_results(keyword, country_code='dk', language_code='lang_da', num_results=5):
#     url = "https://www.googleapis.com/customsearch/v1"
#     params = {
#         "key": st.secrets["google_search_api"],
#         "cx": st.secrets["cx_id"],
#         "q": keyword,
#         "num": num_results,
#         "gl": country_code,
#         "lr": language_code
#     }
#     try:
#         response = requests.get(url, params=params)
#         results = response.json()
#         return [item['link'] for item in results.get("items", []) if "link" in item]
#     except Exception as e:
#         st.warning(f"Error in Google Search API: {e}")
#         return []

# def extract_texts_from_urls(url_list):
#     return " ".join([extract_text_from_url(url) for url in url_list if url])

# def extract_clean_text(documents):
#     return " ".join([doc.page_content for doc in documents])

# def generate_seo_content(text, keyword):
#     text_splitter = RecursiveCharacterTextSplitter(separators=['\n\n', '\n', '.', ','], chunk_size=1000)
#     data = [Document(page_content=text)]
#     docs = text_splitter.split_documents(data)
#     vectorstore = FAISS.from_documents(docs, embeddings)
#     vec = encoder.encode(keyword)
#     retriever = vectorstore.as_retriever(score_threshold=0.7)
#     rdocs = retriever.invoke(keyword)
#     cleaned_text = extract_clean_text(rdocs)
#     query = f"Generate optimized SEO content for {brand} in {language} considering the country : ({country}) which can rank us on No.1. Include: Title (strict character limit of 50-60 characters), Meta (strict character limit of 150-160 characters), H1 (within 60 characters), Intro (strict character limit of 450-500 characters), Info (within 800-1200 words). Avoid competitors brand name in the content."
#     guideline = f"The blog should be in informational and conversational tone for {brand}'s website."
#     guideline_for_seo=f"""Use a primary keyword: Choose one primary keyword to focus on and optimize content around. 
# Include secondary keywords: In addition to primary keyword, include secondary keywords to improve SEO rankings. 
# Use keywords in multiple places: Include keywords in title, meta description, headers, subheadings, and throughout content.
# Include long-tail keywords: use long-tail keywords in the content 
# Synonyms & LSI Keywords, Add related terms and long-tail variations to improve semantic relevance and ranking for multiple queries.
# ⁠Density: Use the main keyword naturally 2–3 times per 100 words (1–3%). Don’t overuse—avoid keyword stuffing.
# Proximity: Keep keywords and their related terms close together in sentences to improve relevance.
# Focus on natural flow and user readability while maintaining SEO signals.
# Avoid keyword stuffing. and duplicate content"""
#     guideline2 = f"You should pick up the activities, events, places, popular attractions for vacation and generic facts about location from the {cleaned_text}. Write it as a human would write and don't use complex words. Make it easier to read and the content should be in proper flow and SEO Optimized."
#     final_query=query + guideline + guideline_for_seo + guideline2 + additional_input
#     response = model.generate_content(final_query)

#     # response = ai_client.invoke(query + guideline + guideline_for_seo + guideline2 + additional_input)
#     return response.text

# # ✅ Main action
# if st.sidebar.button("🚀 Generate SEO Content"):
#     if not brand or not base_urls or not keyword:
#         st.sidebar.error("Please fill in all required fields: Brand, Base URLs, and Keyword.")
#     else:
#         progress_bar = st.progress(0)
#         status_text = st.empty()
#         doc = DC()

#         total_steps = len(base_urls) * 4  # 4 steps per URL
#         current_step = 0

#         for base_url in base_urls:
#             st.write(f"🔗 Base URL: `{base_url}`")
            
#             # Step 1: Extract base URL content
#             status_text.text("Step 1/4: 📄 Extracting content from base URL...")
#             base_text = extract_text_from_url(base_url)
#             current_step += 1
#             progress_bar.progress(current_step / total_steps)

#             # Step 2: Generate keyword
#             status_text.text("Step 2/4: Generating keyword...")
#             # keyword = generate_keyword(base_text)  # You can use AI to generate keyword if not user-provided
#             st.write(f"🔑 Using Keyword: `{keyword}`")
#             current_step += 1
#             progress_bar.progress(current_step / total_steps)

#             # Step 3: Scrape top Google results
#             status_text.text("Step 3/4: 🔍 Scraping competitor URLs...")
#             top_links = scrape_google_search_results(keyword, country_code=search_country, language_code=search_language)
#             st.write("🏁 Top competitor URLs found:")
#             for link in top_links:
#                 st.write(f"- {link}")
#             competitor_text = extract_texts_from_urls(top_links)
#             current_step += 1
#             progress_bar.progress(current_step / total_steps)

#             # Step 4: Generate SEO content
#             status_text.text("Step 4/4: ✍️ Generating SEO content...")
#             generated_content = generate_seo_content(competitor_text, keyword)
#             st.text_area("Generated SEO Content", generated_content, height=300)
#             current_step += 1
#             progress_bar.progress(current_step / total_steps)

#             doc.add_paragraph(f"URL: {base_url.strip()}\n{generated_content}")

#         doc.save(doc_file)
#         st.success("✅ SEO Content generation complete and document saved!")

#         with open(doc_file, "rb") as file:
#             st.download_button(
#                 label="📄 Download Generated Document",
#                 data=file,
#                 file_name="seo_content_revamp.docx",
#                 mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
#             )



# seo_app.py

import os
import io
import time
import requests
import streamlit as st
from bs4 import BeautifulSoup

# LangChain / Vector search
# from langchain.schema import Document
   # Old (deprecated)
   # from langchain.schema import Document

   # New (recommended)
from langchain_core.documents import Document 
# from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings

# Sentence-Transformers (manual, stable init)
from sentence_transformers import SentenceTransformer, models

# Output docx
from docx import Document as DC

# Google Gemini
import google.generativeai as genai


# =========================
# Streamlit UI
# =========================
st.set_page_config(page_title="🧠 SEO Content Generator", page_icon="🤓", layout="wide")

st.markdown(
    """
    <h1 style='text-align: center; color: #4CAF50;'>🧠 SEO Content Generator v2.1 (It is fixed now..)</h1>
    <p style='text-align: center; font-size: 18px;'>Generate top-ranking SEO content. Powered by Gemini Pro & Google Search!</p>
    <hr style="border:1px solid #ccc;">
    """,
    unsafe_allow_html=True
)

col1, col2 = st.columns([6, 1])
with col2:
    st.markdown(
        """
        <a href="https://www.linkedin.com/in/abhishek-kumawat-iitd/" target="_blank">
            <button style="background-color:#0072b1; color:white; padding:8px 16px; border:none; border-radius:6px; cursor:pointer;">
                💡 How it works?
            </button>
        </a>
        """,
        unsafe_allow_html=True
    )

# =========================
# Sidebar
# =========================
st.sidebar.header("🛠️ Input Parameters")
doc_file = "seo_content_revamp_AK.docx"

language_options = {
    "Danish": "lang_da",
    "Dutch": "lang_nl",
    "German": "lang_de",
    "French": "lang_fr",
    "English": "lang_en",
    "Swedish": "lang_sv",
    "Norwegian": "lang_no",
    "Spanish": "lang_es",
}
country_options = {
    "India": "in",
    "Denmark": "dk",
    "Netherlands": "nl",
    "Belgium": "be",
    "Germany": "de",
    "France": "fr",
    "United States": "us",
    "Sweden": "se",
    "Norway": "no",
    "Spain": "es"
}

language_choice = st.sidebar.selectbox("🌐 Language", list(language_options.keys()))
country_choice = st.sidebar.selectbox("🏳️ Country", list(country_options.keys()))

search_language = language_options[language_choice]
search_country = country_options[country_choice]

language = language_choice
country = country_choice
brand = st.sidebar.text_input("🏷️ Brand", "Danland")
base_urls_raw = st.sidebar.text_area("🔗 Base URLs (comma-separated)")
base_urls = [u.strip() for u in base_urls_raw.split(",") if u.strip()]
keyword = st.sidebar.text_area("🔍 Enter the Keyword (Required)")
additional_input = st.sidebar.text_area("✏️ Additional Input (Enter additional input here)")

# =========================
# Config / Secrets
# =========================
DEVICE = "cpu"

# Guardrails for secrets
missing_secrets = []
for key in ("gemini_api", "google_search_api", "cx_id"):
    if key not in st.secrets:
        missing_secrets.append(key)
if missing_secrets:
    st.warning(
        "Some secrets are missing: " + ", ".join(missing_secrets) +
        ". Add them to Streamlit secrets for full functionality."
    )

# =========================
# Caching: Models & Clients
# =========================
@st.cache_resource(show_spinner=True)
def load_encoder_and_embeddings():
    """
    Load SentenceTransformer encoder via module composition to avoid
    meta-tensor issues, and HuggingFaceEmbeddings for FAISS.
    Everything is pinned to CPU.
    """
    # Build encoder from modules (doesn't pass tokenizer objects directly)
    word_embedding_model = models.Transformer("sentence-transformers/all-mpnet-base-v2")
    pooling_model = models.Pooling(word_embedding_model.get_word_embedding_dimension())
    encoder = SentenceTransformer(modules=[word_embedding_model, pooling_model], device=DEVICE)

    # LangChain embeddings (MiniLM) on CPU
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": DEVICE}
    )
    return encoder, embeddings

@st.cache_resource(show_spinner=True)
def load_gemini_client():
    if "gemini_api" in st.secrets:
        genai.configure(api_key=st.secrets["gemini_api"])
        # return genai.GenerativeModel("gemini-2.5-pro")
        return genai.GenerativeModel("gemini-2.5-flash")
        # return genai.GenerativeModel("gemini-1.5-flash")
    return None

encoder, embeddings = load_encoder_and_embeddings()
gemini_model = load_gemini_client()

# =========================
# Helpers
# =========================
def extract_text_from_url(url: str) -> str:
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        soup = BeautifulSoup(response.content, 'html.parser')
        return " ".join(p.get_text(strip=True) for p in soup.find_all("p"))
    except requests.exceptions.RequestException as e:
        st.warning(f"Error fetching {url}: {e}")
        return ""

def scrape_google_search_results(keyword: str, country_code='dk', language_code='lang_da', num_results=5):
    if "google_search_api" not in st.secrets or "cx_id" not in st.secrets:
        st.warning("Google Search API not configured. Provide 'google_search_api' and 'cx_id' in secrets.")
        return []
    url = "https://www.googleapis.com/customsearch/v1"
    params = {
        "key": st.secrets["google_search_api"],
        "cx": st.secrets["cx_id"],
        "q": keyword,
        "num": num_results,
        "gl": country_code,
        "lr": language_code
    }
    try:
        response = requests.get(url, params=params, timeout=15)
        response.raise_for_status()
        results = response.json()
        return [item['link'] for item in results.get("items", []) if "link" in item]
    except Exception as e:
        st.warning(f"Error in Google Search API: {e}")
        return []

def extract_texts_from_urls(url_list):
    texts = []
    for url in url_list:
        if not url:
            continue
        texts.append(extract_text_from_url(url))
    return " ".join(texts)

def extract_clean_text(documents):
    return " ".join([doc.page_content for doc in documents])

def build_vectorstore_from_text(text: str):
    text_splitter = RecursiveCharacterTextSplitter(
        separators=['\n\n', '\n', '.', ','],
        chunk_size=1000,
        chunk_overlap=100
    )
    data = [Document(page_content=text)]
    docs = text_splitter.split_documents(data)
    if not docs:
        return None
    return FAISS.from_documents(docs, embeddings)

def generate_seo_content(cleaned_context: str, keyword: str) -> str:
    if gemini_model is None:
        return "Gemini API is not configured. Please add 'gemini_api' to Streamlit secrets."

    query = (
        f"Generate optimized SEO content for {brand} in {language} considering the country: ({country}) "
        f"which can rank us No.1.\n\n"
        f"Include:\n"
        f"- Title (strict character limit 50–60)\n"
        f"- Meta (strict character limit 150–160)\n"
        f"- H1 (within 60)\n"
        f"- Intro (strict character limit 450–500)\n"
        f"- Info (800–1200 words)\n\n"
        f"Avoid competitor brand names.\n\n"
        f"Context from competitor pages and research:\n{cleaned_context}\n\n"
    )

    guideline = f"The blog should be in informational and conversational tone for {brand}'s website.\n"
    guideline_for_seo = (
        "SEO Guidelines:\n"
        "- Choose one primary keyword and optimize around it.\n"
        "- Include relevant secondary keywords.\n"
        "- Place keywords in title, meta, headers, and naturally in content.\n"
        "- Include long-tail variations and semantically related terms.\n"
        "- Use the main keyword ~1–3% density. Avoid stuffing.\n"
        "- Keep related terms close to the main keyword.\n"
        "- Prioritize readability and natural flow.\n"
    )
    guideline2 = (
        "Use the provided context to pick activities, events, places, popular attractions, and relevant facts. "
        "Write simply (no complex words), maintain a natural flow, and ensure SEO optimization.\n"
    )

    final_query = query + guideline + guideline_for_seo + guideline2 + (additional_input or "")

    try:
        response = gemini_model.generate_content(final_query)
        return response.text
    except Exception as e:
        return f"Error generating content: {e}"

# =========================
# Main Action
# =========================
if st.sidebar.button("🚀 Generate SEO Content"):
    if not brand or not base_urls or not keyword:
        st.sidebar.error("Please fill in all required fields: Brand, Base URLs, and Keyword.")
    else:
        progress_bar = st.progress(0)
        status_text = st.empty()
        doc = DC()

        total_steps = max(1, len(base_urls)) * 4  # 4 steps per URL
        current_step = 0

        for base_url in base_urls:
            st.write(f"🔗 Base URL: `{base_url}`")

            # Step 1: Extract base URL content
            status_text.text("Step 1/4: 📄 Extracting content from base URL...")
            base_text = extract_text_from_url(base_url)
            current_step += 1
            progress_bar.progress(min(1.0, current_step / total_steps))

            # Step 2: Keyword (user-provided; could auto-generate if desired)
            status_text.text("Step 2/4: Generating/using keyword...")
            st.write(f"🔑 Using Keyword: `{keyword}`")
            current_step += 1
            progress_bar.progress(min(1.0, current_step / total_steps))

            # Step 3: Competitor URLs from Google
            status_text.text("Step 3/4: 🔍 Scraping competitor URLs...")
            top_links = scrape_google_search_results(
                keyword, country_code=search_country, language_code=search_language
            )
            if top_links:
                st.write("🏁 Top competitor URLs found:")
                for link in top_links:
                    st.write(f"- {link}")
            else:
                st.info("No competitor URLs found with the current API settings or query.")
            competitor_text = extract_texts_from_urls(top_links)
            # Combine with base text to give more context
            combined_text = (base_text or "") + "\n" + (competitor_text or "")
            current_step += 1
            progress_bar.progress(min(1.0, current_step / total_steps))

            # Step 4: Generate SEO content
            status_text.text("Step 4/4: ✍️ Generating SEO content...")

            # Build vectorstore (optional retrieval to shrink context)
            cleaned_context = ""
            if combined_text.strip():
                vs = build_vectorstore_from_text(combined_text)
                if vs:
                    # Use similarity_search_with_score and filter by score threshold (lower is more similar)
                    hits = vs.similarity_search_with_score(keyword, k=10)
                    # FAISS returns distances; smaller = better. Keep top chunks with distance below a loose cap.
                    # If your FAISS wrapper returns cosine similarity instead, adjust logic accordingly.
                    kept = []
                    for doc, score in hits:
                        kept.append(doc.page_content)
                    cleaned_context = "\n".join(kept)[:8000]  # keep context bounded
                else:
                    cleaned_context = combined_text[:8000]
            else:
                cleaned_context = ""

            generated_content = generate_seo_content(cleaned_context, keyword)
            st.text_area("Generated SEO Content", generated_content, height=350)
            current_step += 1
            progress_bar.progress(min(1.0, current_step / total_steps))

            doc.add_paragraph(f"URL: {base_url.strip()}\n{generated_content}\n")

        # Save and allow download
        doc.save(doc_file)
        st.success("✅ SEO Content generation complete and document saved!")

        with open(doc_file, "rb") as file:
            st.download_button(
                label="📄 Download Generated Document",
                data=file,
                file_name="seo_content_revamp.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

